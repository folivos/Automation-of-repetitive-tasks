#!/usr/bin/env python
# -*- coding: utf-8 -*-
# conexión con trello
from trello import TrelloClient
client = TrelloClient(
    api_key='27d6538f5b3dce7250b9f895955027b6',
    api_secret='bc6b4e855f3e4b16bcf0dfb33c39a6a780d8c6c1f1d6f751c04af442710211fc',
    token='99921e67e99038cde6ed32d48734fef8ed264f1263e7af826397a7776b19df9e',
    token_secret='bc6b4e855f3e4b16bcf0dfb33c39a6a780d8c6c1f1d6f751c04af442710211fc'
)
id_tablero = "5a1392624dbb0d3f446368bd"
tablero = client.get_board(id_tablero) # tablero del informe semanal

# Extraccion de la informacion del tablero Trello
lst = {}
lst["nombre_lista"] = []
lst["info_tarjeta"] = []
lst["nombre_tarjeta"] = []

for lista in tablero.all_lists():
    for tarjeta in lista.list_cards():
        for comentario in tarjeta.get_comments():
            lst["nombre_lista"].append(lista.name)
            lst["nombre_tarjeta"].append(tarjeta.name)
            datos = comentario['data']
            lst["info_tarjeta"].append(datos['text'])

proyectos = lst["nombre_lista"]
datos = lst["info_tarjeta"]
areas = lst["nombre_tarjeta"]

# ordena los datos extraidos
from ordena_informe import Ordena_informe

oMatriz = Ordena_informe()

k =3 # areas x proyecto
matriz_proyecto = []

for i in range(len(oMatriz.vectores(proyectos, areas, datos)[0])):
    aux =[]
    aux.append(oMatriz.vectores(proyectos, areas, datos)[0][i])#proyecto
    for j in range(k-3,k):
        aux.append(oMatriz.vectores(proyectos, areas, datos)[1][j])#area
        aux.append(oMatriz.vectores(proyectos, areas, datos)[2][j])#informacion
    k=k+3
    matriz_proyecto.append(aux)

# Clasificación de Proyectos
from Clasifica_Proyecto import clasifica_proyecto
oClasificador = clasifica_proyecto()

proyectosSTA = oClasificador.clasifica('STA', matriz_proyecto)
proyectosSTN = oClasificador.clasifica('STN', matriz_proyecto)
proyectosPST = oClasificador.clasifica('PST', matriz_proyecto)
proyectosPSA = oClasificador.clasifica('PSA', matriz_proyecto)
proyectosEstudio = oClasificador.clasifica('Estudio', matriz_proyecto)

# creacion y llenado de word: Proyectos en ejecución

import docx
doc = docx.Document('template.docx')

from escribe_parrafo import Escribe_word
oEscribe = Escribe_word()
# Comenzamos a escribir los proyectos clasificados por tipo:

doc.add_paragraph(u"INFORME PROYECTOS EN DESARROLLO", 'Heading 1')
doc.add_paragraph(u"SISTEMA NACIONAL", 'Heading 2')
doc.add_paragraph(u"OBRAS NUEVAS", 'Heading 3')
oEscribe.desarrollo(doc, proyectosSTN)

doc.add_paragraph(u"OBRAS DE AMPLIACIÓN", 'Heading 3')
oEscribe.desarrollo(doc, proyectosSTA)

if len(proyectosPST) == 0:
    pass
else:
    doc.add_paragraph(u"SISTEMA ZONAL", 'Heading 2')
    doc.add_paragraph(u"OBRAS DE AMPLIACIÓN", 'Heading 3')
    oEscribe.desarrollo(doc, proyectosPST)

if len(proyectosPSA) == 0:
    pass
else:
    doc.add_paragraph(u"SISTEMA DEDICADO", 'Heading 2')
    oEscribe.desarrollo(doc, proyectosPSA)

# Creación y llenado de word: Proyectos en Estudio

doc.add_paragraph(u"INFORME PROYECTOS EN ESTUDIO/LICITACIÓN", 'Heading 1')
doc.add_paragraph(u"SISTEMA NACIONAL", 'Heading 2')
doc.add_paragraph(u"OBRAS NUEVAS", 'Heading 3')
# doc.add_paragraph(u"ESTUDIOS", 'Subtitle')

# oEscribe.estudios(doc, proyectosEstudio)




# Guarda archivo word
doc.save('P-PPP-IS-I-2018-08 Informe Semanal (23-feb-2018)_Rev0.docx')